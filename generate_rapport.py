# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ---- PAGE SETUP ----
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ---- STYLES ----
def set_font(run, bold=False, size=12, color=None):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def para_spacing(p, before=0, after=0, line=276):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    from docx.shared import Pt as PT
    from docx.oxml.ns import qn as QN
    from docx.oxml import OxmlElement as OE
    pPr = p._p.get_or_add_pPr()
    lSpc = OE('w:spacing')
    lSpc.set(QN('w:line'), str(line))
    lSpc.set(QN('w:lineRule'), 'auto')
    pPr.append(lSpc)

def add_para(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, size=12, before=0, after=6):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    set_font(run, bold=bold, size=size)
    para_spacing(p, before=before, after=after, line=276)  # 276 = 1.5 line
    return p

def add_heading(text, level=1, num=""):
    full = f"{num} {text}".strip() if num else text
    if level == 0:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(full.upper())
        set_font(run, bold=True, size=14)
        para_spacing(p, before=12, after=6)
    elif level == 1:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(full.upper())
        set_font(run, bold=True, size=13)
        para_spacing(p, before=12, after=6)
    elif level == 2:
        p = doc.add_paragraph()
        run = p.add_run(full.upper())
        set_font(run, bold=True, size=12)
        para_spacing(p, before=10, after=4)
    else:
        p = doc.add_paragraph()
        run = p.add_run(full)
        set_font(run, bold=True, size=12)
        para_spacing(p, before=8, after=4)
    return p

def add_page_break():
    doc.add_page_break()

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run, size=12)
    para_spacing(p, before=0, after=4)

def add_numbered(text, num):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(f"{num}. {text}")
    set_font(run, size=12)
    para_spacing(p, before=0, after=4)

# ======================================================
# PAGE DE GARDE
# ======================================================
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("RÉPUBLIQUE DE CÔTE D'IVOIRE\nUnion – Discipline – Travail")
set_font(r, bold=True, size=12)
para_spacing(p, before=0, after=6)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("[NOM DE L'ÉCOLE / UNIVERSITÉ]")
set_font(r, bold=True, size=13)
para_spacing(p, before=0, after=4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Filière : Informatique Développeur d'Applications (IDA)")
set_font(r, bold=False, size=12)
para_spacing(p, before=0, after=30)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("RAPPORT DE STAGE DE FIN DE FORMATION")
set_font(r, bold=True, size=16, color=(0, 0, 128))
para_spacing(p, before=0, after=10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Thème :")
set_font(r, bold=True, size=13)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CONCEPTION ET RÉALISATION D'UNE APPLICATION DE\nGESTION COMPTABLE : COMPTAFLOW")
set_font(r, bold=True, size=15, color=(180, 0, 0))
para_spacing(p, before=0, after=30)

tbl = doc.add_table(rows=2, cols=2)
tbl.style = 'Table Grid'
cells = [
    ("Réalisé par :", "[Votre Nom et Prénom]"),
    ("Maître de stage :", "Mme AFFIA ESTHER CAROLINE MOKAN"),
]
for i, (lbl, val) in enumerate(cells):
    tbl.cell(i, 0).text = lbl
    tbl.cell(i, 1).text = val
    for cell in [tbl.cell(i, 0), tbl.cell(i, 1)]:
        for run in cell.paragraphs[0].runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Structure d'accueil : LEADER WORLD PERFECT SARL\nAbidjan, Cocody Riviera Attoban, Laurier 3")
set_font(r, size=12)
para_spacing(p, before=20, after=20)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Année académique : 2024 – 2025")
set_font(r, bold=True, size=12)

add_page_break()

# ======================================================
# SOMMAIRE
# ======================================================
add_heading("SOMMAIRE", level=0)

toc_items = [
    ("AVANT-PROPOS", 3),
    ("REMERCIEMENTS", 4),
    ("LISTE DES SIGLES ET ABRÉVIATIONS", 5),
    ("INTRODUCTION GÉNÉRALE", 6),
    ("PREMIÈRE PARTIE : CADRE GÉNÉRAL DU STAGE", 7),
    ("CHAPITRE I : PRÉSENTATION GÉNÉRALE DE L'ENTREPRISE", 8),
    ("I. Historique de l'entreprise", 8),
    ("II. Statut juridique et cadre réglementaire", 9),
    ("III. Secteur d'activité", 9),
    ("IV. Missions et vision", 10),
    ("CHAPITRE II : ORGANISATION ET FONCTIONNEMENT", 11),
    ("I. Organisation structurelle et organigramme", 11),
    ("II. Fonctionnement des différents services", 12),
    ("DEUXIÈME PARTIE : ÉTUDE ET CONCEPTION DU SYSTÈME", 14),
    ("CHAPITRE I : ANALYSE DE L'EXISTANT ET CAHIER DES CHARGES", 15),
    ("I. Présentation du projet COMPTAFLOW", 15),
    ("II. Problématique et objectifs", 16),
    ("III. Spécifications fonctionnelles et techniques", 17),
    ("CHAPITRE II : MODÉLISATION DU SYSTÈME (MERISE)", 18),
    ("I. Règles de gestion", 18),
    ("II. Dictionnaire des données", 19),
    ("III. Structure d'Accès Théorique (SAT)", 21),
    ("IV. Modèle Conceptuel des Données (MCD)", 22),
    ("V. Modèle Logique des Données (MLD)", 24),
    ("TROISIÈME PARTIE : RÉALISATION ET MISE EN ŒUVRE", 25),
    ("CHAPITRE I : ENVIRONNEMENT TECHNIQUE", 25),
    ("I. Outils et technologies utilisés", 25),
    ("II. Architecture de l'application", 26),
    ("CHAPITRE II : PRÉSENTATION DE L'APPLICATION", 27),
    ("I. Tableau de bord et interfaces", 27),
    ("II. Intégration de l'IA pour le scan de factures", 28),
    ("III. État d'avancement et difficultés rencontrées", 29),
    ("CONCLUSION GÉNÉRALE", 30),
    ("BIBLIOGRAPHIE / WEBOGRAPHIE", 31),
    ("ANNEXES", 32),
]

for title, page in toc_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run(title)
    set_font(r1, size=12, bold=("PARTIE" in title or "CHAPITRE" in title or title in ["AVANT-PROPOS","REMERCIEMENTS","LISTE DES SIGLES ET ABRÉVIATIONS","INTRODUCTION GÉNÉRALE","CONCLUSION GÉNÉRALE","BIBLIOGRAPHIE / WEBOGRAPHIE","ANNEXES"]))
    # Add tab fill
    r2 = p.add_run(f" {'.' * max(1, 70 - len(title))} {page}")
    set_font(r2, size=12)
    para_spacing(p, before=0, after=3)

add_page_break()

# ======================================================
# AVANT-PROPOS
# ======================================================
add_heading("AVANT-PROPOS", level=0)
add_para(
    "Ce rapport a été rédigé dans le cadre de notre stage de fin de formation en Informatique Développeur "
    "d'Applications (IDA). Il rend compte de notre expérience au sein de la société LEADER WORLD PERFECT SARL, "
    "où nous avons participé activement au développement d'une solution logicielle de gestion comptable nommée "
    "COMPTAFLOW. Ce stage nous a permis de mettre en pratique les connaissances théoriques acquises durant notre "
    "cursus académique tout en développant des compétences professionnelles indispensables."
)
add_page_break()

# ======================================================
# REMERCIEMENTS
# ======================================================
add_heading("REMERCIEMENTS", level=0)
add_para(
    "Nous adressons nos sincères remerciements à :"
)
add_bullet("Madame AFFIA ESTHER CAROLINE MOKAN, Gérante de LEADER WORLD PERFECT SARL, pour nous avoir accordé cette opportunité de stage.")
add_bullet("L'ensemble du personnel de l'entreprise pour leur accueil chaleureux et leur disponibilité.")
add_bullet("Nos encadreurs pédagogiques pour leurs orientations et conseils tout au long de cette formation.")
add_bullet("Nos familles pour leur soutien indéfectible.")
add_page_break()

# ======================================================
# SIGLES
# ======================================================
add_heading("LISTE DES SIGLES ET ABRÉVIATIONS", level=0)
sigles = [
    ("IDA", "Informatique Développeur d'Applications"),
    ("SARL", "Société à Responsabilité Limitée"),
    ("OHADA", "Organisation pour l'Harmonisation en Afrique du Droit des Affaires"),
    ("SYSCOHADA", "Système Comptable OHADA"),
    ("MCD", "Modèle Conceptuel des Données"),
    ("MLD", "Modèle Logique des Données"),
    ("SAT", "Structure d'Accès Théorique"),
    ("MERISE", "Méthode d'Étude et de Réalisation Informatique pour les Systèmes d'Entreprise"),
    ("MVC", "Modèle-Vue-Contrôleur"),
    ("IA", "Intelligence Artificielle"),
    ("PME", "Petite et Moyenne Entreprise"),
    ("TVA", "Taxe sur la Valeur Ajoutée"),
    ("SGBD", "Système de Gestion de Base de Données"),
    ("PHP", "Hypertext Preprocessor"),
    ("SQL", "Structured Query Language"),
]
tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'
tbl.cell(0,0).text = "Sigle"
tbl.cell(0,1).text = "Signification"
for cell in tbl.rows[0].cells:
    for r in cell.paragraphs[0].runs:
        r.font.name = 'Times New Roman'
        r.font.bold = True
        r.font.size = Pt(12)
for sig, meaning in sigles:
    row = tbl.add_row()
    row.cells[0].text = sig
    row.cells[1].text = meaning
    for cell in row.cells:
        for r in cell.paragraphs[0].runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
add_page_break()

# ======================================================
# INTRODUCTION
# ======================================================
add_heading("INTRODUCTION GÉNÉRALE", level=0)
add_para(
    "Dans un monde en constante mutation numérique, la maîtrise des outils informatiques est devenue un atout "
    "indispensable pour les entreprises, qu'elles soient grandes ou petites. La comptabilité, pilier fondamental "
    "de toute organisation, ne fait pas exception à cette tendance. En Côte d'Ivoire, de nombreuses entreprises "
    "cherchent à moderniser leurs processus de gestion financière pour gagner en efficacité et en fiabilité."
)
add_para(
    "C'est dans ce contexte que s'inscrit notre stage au sein de LEADER WORLD PERFECT SARL, une société "
    "ivoirienne spécialisée dans les services informatiques. La mission principale qui nous a été confiée "
    "consiste à concevoir et réaliser COMPTAFLOW, une application web de gestion comptable conforme aux normes "
    "SYSCOHADA en vigueur en Afrique de l'Ouest."
)
add_para(
    "Ce rapport est structuré en trois grandes parties. La première partie présente le cadre général du stage, "
    "notamment l'entreprise d'accueil et son organisation. La deuxième partie traite de l'étude et de la "
    "conception du système d'information, en s'appuyant sur la méthode MERISE. Enfin, la troisième partie "
    "décrit la phase de réalisation et les technologies mises en œuvre."
)
add_page_break()

# ======================================================
# PARTIE 1
# ======================================================
add_heading("PREMIÈRE PARTIE : CADRE GÉNÉRAL DU STAGE", level=0)
add_page_break()

add_heading("CHAPITRE I : PRÉSENTATION GÉNÉRALE DE L'ENTREPRISE", level=1)
add_heading("I. Historique de l'entreprise", level=2)
add_para(
    "La société LEADER WORLD PERFECT SARL est une entreprise ivoirienne créée en 2021. Constituée sous la "
    "forme d'une Société à Responsabilité Limitée (SARL), elle évolue dans un environnement marqué par la "
    "transformation numérique et la modernisation des systèmes d'information."
)
add_para(
    "Dès sa création, l'entreprise s'est fixée pour objectif d'accompagner les particuliers, les PME et les "
    "grandes entreprises dans la mise en place de solutions informatiques adaptées à leurs besoins. Grâce à "
    "une équipe jeune, dynamique et qualifiée, LEADER WORLD PERFECT SARL s'est progressivement positionnée "
    "comme un acteur émergent dans le domaine des services informatiques en Côte d'Ivoire. Son développement "
    "repose sur trois piliers fondamentaux : l'innovation, la qualité de service et la satisfaction client."
)

add_heading("II. Statut juridique et cadre réglementaire", level=2)
add_para(
    "LEADER WORLD PERFECT SARL est régie par le droit ivoirien, conformément aux dispositions de "
    "l'Organisation pour l'Harmonisation en Afrique du Droit des Affaires (OHADA)."
)
add_para("En tant que SARL, la société présente les caractéristiques suivantes :")
add_bullet("Elle dispose d'une personnalité juridique distincte de celle de ses associés ;")
add_bullet("La responsabilité des associés est limitée à leurs apports ;")
add_bullet("Elle est soumise aux obligations comptables, fiscales et sociales prévues par la législation ivoirienne.")
add_para(
    "Sur le plan fiscal, l'entreprise relève du régime fiscal ivoirien et doit se conformer aux exigences "
    "en matière de déclarations fiscales et de tenue régulière de comptabilité. "
    "Son siège social est situé à : Abidjan, Cocody Riviera Attoban, Laurier 3. Cette localisation "
    "stratégique lui permet d'être proche de sa clientèle et de ses partenaires."
)

add_heading("III. Secteur d'activité", level=2)
add_para(
    "LEADER WORLD PERFECT SARL opère principalement dans le secteur des services informatiques. "
    "Ses activités sont diversifiées et structurées autour de trois principaux pôles :"
)
add_heading("1. Les services informatiques", level=3)
add_para("L'entreprise propose :")
add_bullet("Le développement de sites web professionnels ;")
add_bullet("Le développement d'applications web et mobiles ;")
add_bullet("L'infogérance et l'administration de systèmes informatiques ;")
add_bullet("Le support technique et l'assistance aux utilisateurs.")
add_para(
    "Ce pôle constitue le cœur de métier de l'entreprise et participe activement à la digitalisation des "
    "structures clientes."
)

add_heading("2. La vente et la maintenance de matériel informatique", level=3)
add_para("L'entreprise assure :")
add_bullet("La commercialisation d'équipements informatiques (ordinateurs, périphériques, accessoires) ;")
add_bullet("L'installation et la configuration des équipements ;")
add_bullet("La maintenance préventive et corrective ;")
add_bullet("Le diagnostic et la réparation de matériel.")

add_heading("3. La monétique", level=3)
add_para("LEADER WORLD PERFECT SARL intervient également dans le domaine de la monétique, notamment :")
add_bullet("L'intégration de solutions de paiement électronique ;")
add_bullet("La gestion des systèmes Visa et Mastercard ;")
add_bullet("L'assistance technique liée aux transactions électroniques.")

add_heading("IV. Missions et vision", level=2)
add_para(
    "La vision de LEADER WORLD PERFECT SARL est de devenir le partenaire technologique de référence pour "
    "les entreprises et institutions de Côte d'Ivoire. Ses missions s'articulent autour de :"
)
add_bullet("La fourniture de solutions informatiques innovantes et adaptées aux besoins locaux ;")
add_bullet("L'accompagnement des entreprises dans leur transformation numérique ;")
add_bullet("La maintenance et l'optimisation des systèmes d'information des clients ;")
add_bullet("La formation et la sensibilisation aux bonnes pratiques numériques.")
add_page_break()

add_heading("CHAPITRE II : ORGANISATION ET FONCTIONNEMENT DE L'ENTREPRISE", level=1)
add_heading("I. Organisation structurelle et organigramme", level=2)
add_para(
    "LEADER WORLD PERFECT SARL dispose d'une organisation hiérarchique simple favorisant la fluidité "
    "des échanges et la rapidité des prises de décision. L'organigramme ci-dessous illustre la structure "
    "de l'entreprise :"
)

# Organigramme en tableau stylisé
add_para("[INSÉRER ICI L'ORGANIGRAMME - Voir description ci-dessous]", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(
    "L'organigramme de LEADER WORLD PERFECT SARL se présente de manière hiérarchique avec en tête "
    "La Gérante, dont dépendent directement : l'Assistante de Direction, le Service Informatique, "
    "le Service Technique, le Service Commercial et le Service Comptable.",
    align=WD_ALIGN_PARAGRAPH.CENTER
)

# Draw organigramme as table
org_tbl = doc.add_table(rows=3, cols=5)
org_tbl.style = 'Table Grid'
# Row 0: Gérante (merged)
org_tbl.cell(0, 0).merge(org_tbl.cell(0, 4))
org_tbl.cell(0, 0).text = "La Gérante\nMme AFFIA ESTHER CAROLINE MOKAN"
org_tbl.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in org_tbl.cell(0, 0).paragraphs[0].runs:
    run.font.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
# Row 1: Assistante (merged)
org_tbl.cell(1, 0).merge(org_tbl.cell(1, 4))
org_tbl.cell(1, 0).text = "Assistante de Direction"
org_tbl.cell(1, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in org_tbl.cell(1, 0).paragraphs[0].runs:
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
# Row 2: Services
services = ["Service\nInformatique", "Service\nTechnique", "Service\nCommercial", "Service\nComptable", ""]
for i, svc in enumerate(services):
    org_tbl.cell(2, i).text = svc
    org_tbl.cell(2, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in org_tbl.cell(2, i).paragraphs[0].runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)

add_para("Figure 1 : Organigramme de LEADER WORLD PERFECT SARL", align=WD_ALIGN_PARAGRAPH.CENTER, bold=False)

add_heading("II. Fonctionnement des différents services", level=2)
add_heading("1. La Direction Générale", level=3)
add_para(
    "La direction générale est assurée par Madame AFFIA ESTHER CAROLINE MOKAN, Gérante de la société. "
    "Elle est chargée de définir la vision et les objectifs stratégiques de l'entreprise, de superviser "
    "la gestion financière et budgétaire, et de représenter l'entreprise auprès des partenaires et des clients."
)

add_heading("2. L'Assistante de Direction", level=3)
add_para(
    "L'Assistante de Direction assure le soutien administratif de la gérance : gestion de l'agenda, "
    "organisation des réunions, suivi des correspondances et archivage des documents. "
    "Elle constitue un maillon essentiel dans la coordination interne."
)

add_heading("3. Le Service Commercial", level=3)
add_para(
    "Le service commercial prospecte de nouveaux clients, identifie les besoins du marché, élabore "
    "les offres commerciales et fidélise la clientèle existante. Il contribue directement à la croissance "
    "et à la rentabilité de l'entreprise."
)

add_heading("4. Le Service Informatique", level=3)
add_para(
    "Le service informatique analyse les besoins techniques, développe les applications, assure les tests "
    "et la mise en production, et maintient les systèmes. C'est au sein de ce service que s'inscrit le "
    "développement de l'application COMPTAFLOW, solution conçue pour optimiser la gestion comptable "
    "interne de l'entreprise."
)

add_heading("5. Le Service Technique", level=3)
add_para(
    "Le service technique intervient sur les équipements et infrastructures : installation, configuration "
    "du matériel, maintenance préventive et corrective, intervention en cas de panne et formation des utilisateurs."
)

add_heading("6. Le Service Comptable", level=3)
add_para(
    "Le service comptable gère les opérations financières de l'entreprise : enregistrement des écritures "
    "comptables, déclarations fiscales, suivi budgétaire et établissement des états financiers. "
    "Ce service bénéficie de l'appui du cabinet DC-KNOWING CGA."
)
add_page_break()

# ======================================================
# PARTIE 2
# ======================================================
add_heading("DEUXIÈME PARTIE : ÉTUDE ET CONCEPTION DU SYSTÈME", level=0)
add_page_break()

add_heading("CHAPITRE I : ANALYSE DE L'EXISTANT ET CAHIER DES CHARGES", level=1)
add_heading("I. Présentation du projet COMPTAFLOW", level=2)
add_para(
    "COMPTAFLOW est une application web de gestion comptable développée sur le framework Laravel 12. "
    "Elle est conçue pour répondre aux besoins comptables des entreprises ivoiriennes, en conformité "
    "avec le plan comptable SYSCOHADA. L'application adopte une architecture multi-sociétés permettant "
    "à plusieurs entités d'utiliser la même plateforme de manière totalement isolée."
)
add_para(
    "L'une des innovations majeures de COMPTAFLOW est l'intégration de l'Intelligence Artificielle "
    "(via l'API Gemini de Google DeepMind) pour l'extraction automatique des données de factures "
    "numérisées, réduisant considérablement le temps de saisie manuelle."
)

add_heading("II. Problématique et objectifs", level=2)
add_heading("1. Problématique", level=3)
add_para(
    "Avant COMPTAFLOW, la gestion comptable au sein de LEADER WORLD PERFECT SARL s'effectuait à l'aide "
    "de tableurs (Excel) et de fichiers papier. Cette situation engendrait plusieurs problèmes :"
)
add_bullet("Risque élevé d'erreurs de saisie et de calcul ;")
add_bullet("Difficulté de suivi et de contrôle des opérations ;")
add_bullet("Perte de temps lors de la génération des états comptables ;")
add_bullet("Absence de traçabilité des modifications effectuées.")

add_heading("2. Objectifs", level=3)
add_para("Le projet COMPTAFLOW vise à :")
add_bullet("Automatiser la saisie des écritures grâce au scan intelligent de factures ;")
add_bullet("Centraliser les données comptables dans une base de données sécurisée ;")
add_bullet("Générer automatiquement les états comptables (Grand Livre, Balance) ;")
add_bullet("Gérer les droits d'accès par rôle (Administrateur, Comptable) .")

add_heading("III. Spécifications fonctionnelles et techniques", level=2)
add_heading("1. Spécifications fonctionnelles", level=3)
add_para("L'application doit permettre :")
add_bullet("La gestion d'un plan comptable SYSCOHADA (création, modification, suppression) ;")
add_bullet("La gestion des tiers (clients, fournisseurs) ;")
add_bullet("La configuration des journaux comptables (Achats, Ventes, Caisse, Banque) ;")
add_bullet("La saisie et la validation des écritures comptables ;")
add_bullet("La gestion des exercices comptables avec clôture ;")
add_bullet("L'édition du Grand Livre et de la Balance ;")
add_bullet("Le scan et l'interprétation automatique des factures via l'IA.")

add_heading("2. Spécifications techniques", level=3)
add_para("La solution repose sur :")
add_bullet("Framework backend : Laravel 12 (PHP 8.2+) ;")
add_bullet("Base de données : MySQL 8.0 ;")
add_bullet("Frontend : Blade templates + TailwindCSS ;")
add_bullet("IA : API Google Gemini pour l'OCR intelligent ;")
add_bullet("Exports : Laravel Excel, DomPDF.")
add_page_break()

add_heading("CHAPITRE II : MODÉLISATION DU SYSTÈME (MÉTHODE MERISE)", level=1)
add_heading("I. Règles de gestion", level=2)
add_para(
    "Les règles de gestion définissent les contraintes métier fondamentales qui régissent le "
    "fonctionnement de l'application. Elles ont été établies sur la base de l'analyse des processus "
    "comptables de l'entreprise :"
)
regles = [
    ("RG1", "Une société (Company) peut posséder plusieurs utilisateurs."),
    ("RG2", "Un utilisateur appartient à une seule société et dispose d'un rôle unique (Admin ou Comptable)."),
    ("RG3", "Une société possède un plan comptable composé de plusieurs comptes."),
    ("RG4", "Un compte comptable (Plan Comptable) appartient à une seule société."),
    ("RG5", "Un compte tiers (Plan Tiers) est rattaché à un compte du plan comptable général."),
    ("RG6", "Une écriture comptable doit être rattachée à un journal, un exercice, un utilisateur et une société."),
    ("RG7", "Une écriture comptable mouvemente un compte du plan comptable et un compte du plan tiers."),
    ("RG8", "Pour chaque numéro de pièce, la somme des débits doit être égale à la somme des crédits."),
    ("RG9", "Un journal saisi (Journaux Saisis) est créé pour chaque combinaison exercice/journal/mois."),
    ("RG10", "Un exercice comptable clôturé ne peut plus recevoir de nouvelles écritures."),
    ("RG11", "Les droits d'accès d'un utilisateur sont définis par ses habilitations configurées par l'administrateur."),
    ("RG12", "Un Grand Livre est généré à partir des écritures d'un journal sur une période donnée."),
]
for code, texte in regles:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{code} : ")
    set_font(r1, bold=True, size=12)
    r2 = p.add_run(texte)
    set_font(r2, size=12)
    para_spacing(p, before=0, after=4)

add_heading("II. Dictionnaire des données", level=2)
add_para(
    "Le dictionnaire des données recense l'ensemble des informations utilisées dans le système "
    "COMPTAFLOW. Il précise pour chaque donnée son identifiant, son libellé, son type, sa longueur "
    "et sa nature."
)

dd_headers = ["Nom", "Libellé", "Type", "Longueur", "Nature"]
dd_rows = [
    ("id_soc", "Identifiant société", "Entier", "10", "Identifiant"),
    ("nom_soc", "Nom de la société", "Alpha-num.", "191", "Elémentaire"),
    ("forme_jur", "Forme juridique", "Alpha-num.", "191", "Elémentaire"),
    ("cap_soc", "Capital social", "Décimal", "15,2", "Elémentaire"),
    ("adresse_soc", "Adresse siège social", "Alpha-num.", "191", "Elémentaire"),
    ("tel_soc", "Numéro de téléphone", "Alpha-num.", "191", "Elémentaire"),
    ("email_soc", "Adresse e-mail", "Alpha-num.", "191", "Elémentaire"),
    ("id_user", "Identifiant utilisateur", "Entier", "10", "Identifiant"),
    ("nom_user", "Nom de l'utilisateur", "Alpha-num.", "191", "Elémentaire"),
    ("prenom_user", "Prénom de l'utilisateur", "Alpha-num.", "191", "Elémentaire"),
    ("email_user", "E-mail utilisateur", "Alpha-num.", "191", "Elémentaire"),
    ("role_user", "Rôle (admin/comptable)", "Enum.", "-", "Elémentaire"),
    ("habilit", "Habilitations JSON", "JSON", "-", "Calculée"),
    ("id_ex", "Identifiant exercice", "Entier", "10", "Identifiant"),
    ("date_deb", "Date début exercice", "Date", "-", "Elémentaire"),
    ("date_fin", "Date fin exercice", "Date", "-", "Elémentaire"),
    ("cloture", "Exercice clôturé (O/N)", "Booléen", "-", "Elémentaire"),
    ("id_compte", "Identifiant compte", "Entier", "10", "Identifiant"),
    ("num_compte", "Numéro de compte", "Alpha-num.", "191", "Elémentaire"),
    ("intit_compte", "Intitulé du compte", "Alpha-num.", "191", "Elémentaire"),
    ("type_compte", "Type (Bilan/Résultat)", "Alpha-num.", "191", "Elémentaire"),
    ("classe", "Classe SYSCOHADA (1-8)", "Entier", "2", "Elémentaire"),
    ("id_tiers", "Identifiant tiers", "Entier", "10", "Identifiant"),
    ("num_tiers", "Numéro de tiers", "Alpha-num.", "191", "Elémentaire"),
    ("intit_tiers", "Intitulé du tiers", "Alpha-num.", "191", "Elémentaire"),
    ("type_tiers", "Type (Client/Fournisseur)", "Alpha-num.", "191", "Elémentaire"),
    ("id_journal", "Identifiant journal", "Entier", "10", "Identifiant"),
    ("code_journal", "Code du journal (ex: ACH)", "Alpha-num.", "191", "Elémentaire"),
    ("type_journal", "Type journal (Achats, Ventes…)", "Alpha-num.", "191", "Elémentaire"),
    ("id_ecrit", "Identifiant écriture", "Entier", "10", "Identifiant"),
    ("date_ecrit", "Date de l'écriture", "Date", "-", "Elémentaire"),
    ("n_saisie", "Numéro de saisie", "Alpha-num.", "191", "Elémentaire"),
    ("ref_piece", "Référence de la pièce", "Alpha-num.", "191", "Elémentaire"),
    ("libelle", "Description de l'opération", "Alpha-num.", "191", "Elémentaire"),
    ("montant_db", "Montant au débit", "Décimal", "15,2", "Elémentaire"),
    ("montant_cr", "Montant au crédit", "Décimal", "15,2", "Elémentaire"),
    ("pj", "Pièce justificative (fichier)", "Alpha-num.", "191", "Elémentaire"),
]
tbl = doc.add_table(rows=1, cols=5)
tbl.style = 'Table Grid'
for i, h in enumerate(dd_headers):
    cell = tbl.cell(0, i)
    cell.text = h
    for r in cell.paragraphs[0].runs:
        r.font.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
for row_data in dd_rows:
    row = tbl.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val
        for r in row.cells[i].paragraphs[0].runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)

add_para("Tableau 1 : Dictionnaire des données de COMPTAFLOW", align=WD_ALIGN_PARAGRAPH.CENTER, bold=False)
add_page_break()

# ======================================================
# SAT
# ======================================================
add_heading("III. Structure d'Accès Théorique (SAT)", level=2)
add_para(
    "La Structure d'Accès Théorique (SAT) représente les entités du système avec "
    "leurs attributs et les liens d'accès entre elles. L'identifiant de chaque entité est souligné. "
    "Les flèches indiquent les relations d'identification (qui identifie qui) :"
)

sat_lines = [
    "COMPANY",
    "  id_soc  nom_soc  forme_jur  cap_soc  adresse_soc  tel_soc  email_soc",
    "      |",
    "      | (1,n)               (1,1)",
    "      V",
    "USER",
    "  id_user  nom_user  prenom_user  email_user  role_user  habilit  #id_soc",
    "      |",
    "      | (0,n)               (1,1)",
    "      V",
    "EXERCICE_COMPTABLE",
    "  id_ex  date_deb  date_fin  intitule  cloture  #id_soc  #id_user",
    "      |",
    "      | (1,n)               (1,1)",
    "      V",
    "ECRITURE_COMPTABLE",
    "  id_ecrit  date_ecrit  n_saisie  ref_piece  libelle  montant_db  montant_cr  pj",
    "  #id_ex  #id_journal  #id_compte  #id_tiers  #id_user  #id_soc",
    "      |                     |",
    "      | (0,n)         (0,n) |",
    "      V                     V",
    "CODE_JOURNAL             PLAN_COMPTABLE",
    "  id_journal               id_compte",
    "  code_journal              num_compte",
    "  intitule                  intit_compte",
    "  type_journal              type_compte  classe",
    "                                |",
    "                          (1,n) | (1,1)",
    "                                V",
    "                        PLAN_TIERS",
    "                          id_tiers  num_tiers  intit_tiers  type_tiers",
    "                          #id_compte  #id_soc  #id_user",
]

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("\n".join(sat_lines))
r.font.name = 'Courier New'
r.font.size = Pt(9)
para_spacing(p, before=6, after=6, line=240)

add_para("Figure 2 : Structure d'Accès Théorique (SAT) de COMPTAFLOW", align=WD_ALIGN_PARAGRAPH.CENTER, bold=False)
add_page_break()

# ======================================================
# MCD
# ======================================================
add_heading("IV. Modèle Conceptuel des Données (MCD)", level=2)
add_para(
    "Le Modèle Conceptuel des Données (MCD) représente graphiquement les entités du système "
    "et leurs associations, avec les cardinalités précises de chaque côté de chaque lien. "
    "Le schéma ci-dessous illustre les relations entre les entités principales de COMPTAFLOW :"
)
add_para("[INSÉRER ICI LE SCHÉMA MCD - Description ci-dessous]", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

add_para("Description des entités et associations avec cardinalités :", bold=True)

associations = [
    ("COMPANY – USER", "APPARTENIR", "1,n", "1,1",
     "Une société possède un ou plusieurs utilisateurs. Un utilisateur appartient à une seule société."),
    ("COMPANY – PLAN_COMPTABLE", "DÉFINIR", "1,n", "1,1",
     "Une société définit un ou plusieurs comptes comptables. Un compte appartient à une société."),
    ("COMPANY – CODE_JOURNAL", "CONFIGURER", "1,n", "1,1",
     "Une société configure plusieurs journaux. Un journal appartient à une société."),
    ("COMPANY – EXERCICE", "GÉRER", "1,n", "1,1",
     "Une société gère plusieurs exercices. Un exercice est rattaché à une société."),
    ("USER – ECRITURE", "SAISIR", "0,n", "1,1",
     "Un utilisateur peut saisir aucune ou plusieurs écritures. Une écriture est saisie par un seul utilisateur."),
    ("EXERCICE – JOURNAUX_SAISIS", "CONTENIR", "1,n", "1,1",
     "Un exercice contient plusieurs journaux saisis. Un journal saisi appartient à un exercice."),
    ("CODE_JOURNAL – JOURNAUX_SAISIS", "CONCERNER", "1,n", "1,1",
     "Un code journal est associé à plusieurs journaux saisis. Un journal saisi concerne un code journal."),
    ("JOURNAUX_SAISIS – ECRITURE", "REGROUPER", "1,n", "1,1",
     "Un journal saisi regroupe plusieurs écritures. Une écriture appartient à un journal saisi."),
    ("PLAN_COMPTABLE – ECRITURE", "MOUVEMENTER", "0,n", "1,1",
     "Un compte peut être mouvementé par plusieurs écritures. Une écriture mouvemente un seul compte."),
    ("PLAN_TIERS – ECRITURE", "ÊTRE LIÉ À", "0,n", "1,1",
     "Un tiers peut être lié à plusieurs écritures. Une écriture est liée à un seul tiers."),
    ("PLAN_COMPTABLE – PLAN_TIERS", "SUBDIVISER", "0,n", "1,1",
     "Un compte général peut être subdivisé en plusieurs tiers. Un tiers est rattaché à un seul compte général."),
]

tbl = doc.add_table(rows=1, cols=5)
tbl.style = 'Table Grid'
headers = ["Entité 1", "Association", "Card. (E1)", "Card. (E2)", "Entité 2"]
for i, h in enumerate(headers):
    tbl.cell(0, i).text = h
    for r in tbl.cell(0, i).paragraphs[0].runs:
        r.font.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)
    tbl.cell(0, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for e1_e2, assoc, c1, c2, desc in associations:
    row = tbl.add_row()
    parts = e1_e2.split(" – ")
    row.cells[0].text = parts[0]
    row.cells[1].text = assoc
    row.cells[2].text = c1
    row.cells[3].text = c2
    row.cells[4].text = parts[1] if len(parts) > 1 else ""
    for cell in row.cells:
        for r in cell.paragraphs[0].runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)

add_para("Tableau 2 : Associations et cardinalités du MCD de COMPTAFLOW", align=WD_ALIGN_PARAGRAPH.CENTER)
add_page_break()

# ======================================================
# MLD
# ======================================================
add_heading("V. Modèle Logique des Données (MLD)", level=2)
add_para(
    "Le MLD est la traduction du MCD dans un modèle relationnel. Chaque entité devient une table "
    "dont la clé primaire est notée en gras et les clés étrangères sont préfixées par #."
)
mld_tables = [
    "COMPANY(#id_soc, nom_soc, forme_jur, cap_soc, adresse_soc, tel_soc, email_soc)",
    "USER(#id_user, nom_user, prenom_user, email_user, role_user, habilit, id_soc*)",
    "EXERCICE_COMPTABLE(#id_ex, date_deb, date_fin, intitule, cloture, id_soc*, id_user*)",
    "PLAN_COMPTABLE(#id_compte, num_compte, intit_compte, type_compte, classe, id_soc*, id_user*)",
    "PLAN_TIERS(#id_tiers, num_tiers, intit_tiers, type_tiers, id_compte*, id_soc*, id_user*)",
    "CODE_JOURNAL(#id_journal, code_journal, intitule, type_journal, compte_contrep., id_soc*, id_user*)",
    "JOURNAUX_SAISIS(#id_jsaisi, annee, mois, id_ex*, id_journal*, id_soc*, id_user*)",
    "ECRITURE_COMPTABLE(#id_ecrit, date_ecrit, n_saisie, ref_piece, libelle, montant_db, montant_cr, pj, id_jsaisi*, id_compte*, id_tiers*, id_user*, id_soc*)",
    "GRAND_LIVRE(#id_gl, date_deb, date_fin, grand_livre, id_journal*, id_soc*, id_user*)",
]
for t in mld_tables:
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.font.name = 'Courier New'
    r.font.size = Pt(10)
    para_spacing(p, before=2, after=2, line=240)
add_page_break()

# ======================================================
# PARTIE 3
# ======================================================
add_heading("TROISIÈME PARTIE : RÉALISATION ET MISE EN ŒUVRE", level=0)
add_page_break()

add_heading("CHAPITRE I : ENVIRONNEMENT TECHNIQUE", level=1)
add_heading("I. Outils et technologies utilisés", level=2)
add_heading("1. Laravel 12 (Framework Backend)", level=3)
add_para(
    "Laravel est un framework PHP open-source, réputé pour son expressivité et sa robustesse. "
    "Il suit le patron d'architecture MVC (Modèle-Vue-Contrôleur) et offre :"
)
add_bullet("Un système d'authentification et de gestion des sessions sécurisé ;")
add_bullet("Le moteur de template Blade pour des vues dynamiques ;")
add_bullet("Eloquent ORM pour une interaction simplifiée avec la base de données ;")
add_bullet("Les Migrations pour la gestion versionnée du schéma de base de données.")

add_heading("2. MySQL 8.0 (Base de données)", level=3)
add_para(
    "MySQL est le SGBD relationnel retenu pour son excellente performance, sa fiabilité et sa "
    "compatibilité native avec Laravel. La base de données comprend 9 tables principales gérant "
    "les sociétés, utilisateurs, comptes, écritures et états comptables."
)

add_heading("3. TailwindCSS (Interface utilisateur)", level=3)
add_para(
    "TailwindCSS est un framework CSS utilitaire permettant de créer des interfaces modernes, "
    "élégantes et entièrement responsives. Il garantit une expérience utilisateur fluide sur "
    "tous les types d'appareils (PC, tablette, mobile)."
)

add_heading("4. Intelligence Artificielle – Google Gemini", level=3)
add_para(
    "L'API Gemini de Google DeepMind est intégrée pour analyser et extraire automatiquement les "
    "données des factures numérisées. L'utilisateur soumet une image ou un scan de facture, et "
    "l'IA identifie le montant, la date, le fournisseur et le numéro de pièce, pré-remplissant "
    "ainsi le formulaire de saisie. Cette fonctionnalité réduit le temps de saisie d'environ 70%."
)

add_heading("II. Architecture de l'application", level=2)
add_para(
    "COMPTAFLOW suit une architecture en couches basée sur le patron MVC :"
)
add_bullet("Modèles (Models) : Représentent les entités métier (EcritureComptable, PlanComptable, User…) et leurs relations Eloquent ;")
add_bullet("Vues (Views) : Interfaces Blade dynamiques organisées par module (journaux, plan comptable, dashboard…) ;")
add_bullet("Contrôleurs (Controllers) : Orchestrent la logique applicative et font le lien entre les modèles et les vues ;")
add_bullet("Middlewares : Sécurisent les routes et vérifient les habilitations de l'utilisateur connecté.")
add_para(
    "La gestion multi-sociétés est assurée par l'ajout systématique d'un champ company_id dans "
    "chaque table métier, garantissant l'isolation totale des données entre les différentes "
    "sociétés utilisant la plateforme."
)
add_page_break()

add_heading("CHAPITRE II : PRÉSENTATION DE L'APPLICATION", level=1)
add_heading("I. Tableau de bord et interfaces principales", level=2)
add_para(
    "L'interface principale de COMPTAFLOW se présente sous la forme d'un tableau de bord (dashboard) "
    "offrant une vue d'ensemble des activités comptables :"
)
add_bullet("Affichage du solde des comptes de trésorerie en temps réel ;")
add_bullet("Résumé des dernières écritures saisies ;")
add_bullet("Accès rapide aux modules : Journal, Plan Comptable, Plan Tiers, États Financiers.")
add_para("[INSÉRER CAPTURE D'ÉCRAN DU DASHBOARD]", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Figure 3 : Tableau de bord de COMPTAFLOW", align=WD_ALIGN_PARAGRAPH.CENTER)

add_heading("II. Intégration de l'IA pour le scan de factures", level=2)
add_para(
    "La fonctionnalité de scan de factures par IA constitue l'une des innovations majeures de "
    "COMPTAFLOW. Le processus se déroule en trois étapes :"
)
add_numbered("L'utilisateur télécharge l'image ou le PDF de la facture dans l'interface dédiée.", "1")
add_numbered("L'image est transmise à l'API Google Gemini qui analyse le document et extrait les informations clés.", "2")
add_numbered("Les données extraites (montant, date, référence, tiers) sont pré-remplies dans le formulaire de saisie que l'utilisateur n'a plus qu'à valider.", "3")
add_para("[INSÉRER CAPTURE D'ÉCRAN DU MODULE DE SCAN]", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Figure 4 : Interface de scan de factures assistée par IA", align=WD_ALIGN_PARAGRAPH.CENTER)

add_heading("III. État d'avancement et difficultés rencontrées", level=2)
add_heading("1. État d'avancement", level=3)
add_para("À l'issue du stage, les modules suivants sont fonctionnels :")
add_bullet("Gestion des sociétés et des utilisateurs avec habilitations ;")
add_bullet("Plan comptable SYSCOHADA (import, création manuelle) ;")
add_bullet("Plan tiers (clients et fournisseurs) ;")
add_bullet("Codes journaux (Achats, Ventes, Caisse, Banque) ;")
add_bullet("Saisie des écritures comptables avec contrôle d'équilibre ;")
add_bullet("Gestion des exercices comptables avec clôture ;")
add_bullet("Grand Livre et Balance (consultation et export) ;")
add_bullet("Scan de factures par IA (Gemini).")

add_heading("2. Difficultés rencontrées", level=3)
add_para("Plusieurs défis ont été rencontrés au cours du développement :")
add_bullet("La gestion de la cohérence multi-sociétés nécessitant une rigueur accrue dans toutes les requêtes ;")
add_bullet("L'intégration de l'API Gemini avec gestion des différents formats de factures ;")
add_bullet("La conformité stricte au plan comptable SYSCOHADA avec ses 8 classes et plusieurs centaines de comptes ;")
add_bullet("La conception de l'interface utilisateur pour être à la fois intuitive et fonctionnelle.")
add_page_break()

# ======================================================
# CONCLUSION
# ======================================================
add_heading("CONCLUSION GÉNÉRALE", level=0)
add_para(
    "Ce stage au sein de LEADER WORLD PERFECT SARL a constitué une expérience professionnelle "
    "extrêmement enrichissante. Il nous a permis de mettre en pratique l'ensemble des compétences "
    "acquises tout au long de notre formation en Informatique Développeur d'Applications (IDA)."
)
add_para(
    "La conception et la réalisation de COMPTAFLOW illustrent parfaitement la démarche d'analyse "
    "et de développement logiciel que nous avons apprise. En partant d'une problématique réelle "
    "d'entreprise – l'automatisation de la gestion comptable – nous avons abouti à une solution "
    "opérationnelle intégrant des technologies modernes telles que Laravel, MySQL et l'Intelligence "
    "Artificielle via l'API Google Gemini."
)
add_para(
    "La méthode MERISE nous a guidés efficacement dans la phase de modélisation, garantissant une "
    "conception rigoureuse et évolutive du système d'information. Les cardinalités définies dans le "
    "MCD assurent l'intégrité des données, tandis que la SAT clarifie les accès et les dépendances "
    "entre entités."
)
add_para(
    "Ce projet constitue une base solide et extensible. Des développements futurs pourraient inclure "
    "la gestion complète de la TVA, un module de paie et ressources humaines, ainsi qu'un module "
    "de rapprochement bancaire automatisé. COMPTAFLOW a le potentiel de devenir une référence dans "
    "le domaine de la gestion comptable numérique en Côte d'Ivoire."
)
add_page_break()

# ======================================================
# BIBLIOGRAPHIE
# ======================================================
add_heading("BIBLIOGRAPHIE / WEBOGRAPHIE", level=0)
refs = [
    "OHADA, Plan Comptable Général SYSCOHADA Révisé, édition 2017.",
    "LARAVEL Documentation officielle. https://laravel.com/docs/12.x",
    "GOOGLE DEEPMIND, Gemini API Documentation. https://ai.google.dev/",
    "MERISE, Méthode complète de modélisation. Editions Eyrolles.",
    "PHP Documentation officielle. https://www.php.net/docs.php",
    "MySQL 8.0 Reference Manual. https://dev.mysql.com/doc/",
    "TailwindCSS Documentation. https://tailwindcss.com/docs",
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    r = p.add_run(f"[{i}] {ref}")
    set_font(r, size=12)
    para_spacing(p, before=0, after=6)
add_page_break()

# ======================================================
# ANNEXES
# ======================================================
add_heading("ANNEXES", level=0)
add_heading("Annexe 1 : Schéma de la base de données", level=2)
add_para("Tables principales de la base de données bd_flow_compta :")
tables_db = [
    "companies        : id, company_name, activity, juridique_form, social_capital, adresse, city, phone_number, email_adresse",
    "users            : id, name, last_name, email_adresse, password, role, habilitations, company_id",
    "exercices_comptables : id, date_debut, date_fin, intitule, cloturer, user_id, company_id",
    "plan_comptables  : id, numero_de_compte, intitule, type_de_compte, poste, classe, user_id, company_id",
    "plan_tiers       : id, numero_de_tiers, compte_general, intitule, type_de_tiers, user_id, company_id",
    "code_journals    : id, code_journal, intitule, type, traitement_analytique, user_id, company_id",
    "journaux_saisis  : id, annee, mois, exercices_comptables_id, code_journals_id, user_id, company_id",
    "ecriture_comptables : id, date, n_saisie, description_operation, reference_piece, plan_comptable_id, plan_tiers_id, code_journal_id, exercices_comptables_id, journaux_saisis_id, debit, credit, user_id, company_id",
    "grand_livres     : id, date_debut, date_fin, code_journals_id, grand_livre, user_id, company_id",
]
for t in tables_db:
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.font.name = 'Courier New'
    r.font.size = Pt(9)
    para_spacing(p, before=0, after=2, line=240)

# ======================================================
# SAUVEGARDE
# ======================================================
output_path = r"c:\laragon\www\COMPTAFLOW\rapport-soutenance-FINAL.docx"
doc.save(output_path)
print(f"Rapport sauvegardé : {output_path}")
